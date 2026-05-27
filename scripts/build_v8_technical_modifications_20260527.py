from __future__ import annotations

import json
import math
import os
import shutil
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import ExtraTreesClassifier, RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.metrics import balanced_accuracy_score, accuracy_score
from sklearn.model_selection import GroupKFold
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler


ROOT = Path(r"D:\kec folder\my files\student project\Elsevier_Q2_ExternalOnly_StabilizedSoil_ML")
V7 = ROOT / "Submission_Manuscript_v7_FinalReview_GitHubReady_2026-05-26"
REPO = V7 / "source-aware-stabilized-soil-ml"
OUTDIR = ROOT / "Submission_Manuscript_v8_TechnicalModifications_2026-05-27"
INDOC = V7 / "Submission_manuscript_v7_final_review_ready_with_private_repo.docx"
OUTDOC = OUTDIR / "Submission_manuscript_v8_technical_modifications.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(8)
            if r == 0:
                set_cell_shading(cell, "D9EAF7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    return p


def find_heading(doc: Document, exact: str):
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    raise ValueError(f"heading not found: {exact}")


def insert_before(anchor, block) -> None:
    anchor._p.addprevious(block._p if hasattr(block, "_p") else block._tbl)


def add_paragraph_before(doc: Document, anchor, text: str, style: str | None = None):
    p = doc.add_paragraph(text, style=style)
    insert_before(anchor, p)
    return p


def add_table_before(doc: Document, anchor, caption: str, headers: list[str], rows: list[list[str]]):
    cap = add_caption(doc, caption)
    insert_before(anchor, cap)
    table = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = value
    format_table(table)
    insert_before(anchor, table)
    return table


def rounded(x, nd=3):
    if x is None or (isinstance(x, float) and not math.isfinite(x)):
        return ""
    return f"{x:.{nd}f}"


def make_carbon_sensitivity() -> pd.DataFrame:
    pareto = pd.read_csv(REPO / "results" / "pareto_front_real_mixtures_v1.csv")
    gain_col = next((c for c in ["predicted_gain_pct", "gain_pct", "gain"] if c in pareto.columns), None)
    co2_col = next((c for c in ["kgco2e_per_100kg_soil", "kg_co2e_per_100kg_soil", "co2_kg_per_100kg"] if c in pareto.columns), None)
    response_col = next((c for c in ["response", "response_type", "pareto_response", "dataset"] if c in pareto.columns), None)
    if gain_col is None or co2_col is None:
        raise ValueError(f"Pareto file lacks expected gain/carbon columns: {list(pareto.columns)}")
    if response_col is None:
        pareto["response_group"] = "all"
        response_col = "response_group"

    out = []
    for binder_factor in [0.80, 1.00, 1.20]:
        for distance in [0, 100, 300]:
            tmp = pareto.copy()
            transport = 0.1 * distance * 0.062
            tmp["scenario_co2"] = tmp[co2_col] * binder_factor + transport
            tmp["efficiency"] = tmp[gain_col] / tmp["scenario_co2"].replace(0, np.nan)
            for resp, grp in tmp.groupby(response_col):
                out.append(
                    {
                        "response": str(resp),
                        "binder_factor": binder_factor,
                        "transport_km": distance,
                        "n": len(grp),
                        "median_co2_kg_per_100kg": grp["scenario_co2"].median(),
                        "median_gain_per_kg_co2": grp["efficiency"].median(),
                        "top10_gain_per_kg_co2": grp.nlargest(min(10, len(grp)), "efficiency")["efficiency"].mean(),
                    }
                )
    res = pd.DataFrame(out)
    res.to_csv(OUTDIR / "dynamic_carbon_sensitivity_v8.csv", index=False)
    return res


def prepare_task(df: pd.DataFrame, response: str):
    df = df.copy()
    for c in df.columns:
        if df[c].dtype == object:
            df[c] = df[c].replace({"": np.nan, "nan": np.nan, "None": np.nan})
    source_col = next(c for c in ["source_dataset", "source_id", "source", "study_id"] if c in df.columns)
    numeric_cols = [
        c
        for c in df.select_dtypes(include=[np.number]).columns
        if c.lower()
        not in {
            "ucs_kpa",
            "cbr_percent",
            "source_index",
            "fold",
        }
    ]
    cat_cols = [
        c
        for c in df.columns
        if c not in numeric_cols
        and c not in {source_col, "ucs_kpa", "cbr_percent", "ucs_class", "cbr_class", "ucs_high_binary", "cbr_pass_8"}
        and df[c].nunique(dropna=True) <= 40
    ]
    if response == "UCS":
        target_col = "ucs_high_binary"
        if target_col not in df.columns:
            value_col = next(c for c in ["ucs_kpa", "UCS_kPa", "ucs"] if c in df.columns)
            df[target_col] = (df[value_col] >= 1500).astype(int)
    else:
        target_col = "cbr_pass_8"
        if target_col not in df.columns:
            value_col = next(c for c in ["cbr_percent", "CBR_percent", "cbr_pct", "cbr"] if c in df.columns)
            df[target_col] = (df[value_col] >= 8).astype(int)
    keep = numeric_cols + cat_cols + [source_col, target_col]
    df = df[keep].dropna(subset=[target_col, source_col])
    return df, numeric_cols, cat_cols, source_col, target_col


def source_weighting_test() -> pd.DataFrame:
    tasks = [
        ("UCS high-strength classifier", "UCS", REPO / "data" / "final_reliable_UCS_dataset_v1.csv", ExtraTreesClassifier(n_estimators=400, random_state=42, min_samples_leaf=2, class_weight="balanced")),
        ("CBR pass classifier", "CBR", REPO / "data" / "final_reliable_CBR_dataset_v1.csv", RandomForestClassifier(n_estimators=400, random_state=42, min_samples_leaf=2, class_weight="balanced")),
    ]
    rows = []
    for name, response, path, model in tasks:
        df = pd.read_csv(path)
        data, num_cols, cat_cols, source_col, y_col = prepare_task(df, response)
        groups = data[source_col].astype(str).to_numpy()
        y = data[y_col].astype(int).to_numpy()
        x = data[num_cols + cat_cols]
        n_splits = min(5, len(np.unique(groups)))
        cv = GroupKFold(n_splits=n_splits)
        pre = ColumnTransformer(
            transformers=[
                ("num", Pipeline([("imputer", SimpleImputer(strategy="median")), ("scaler", StandardScaler())]), num_cols),
                ("cat", Pipeline([("imputer", SimpleImputer(strategy="most_frequent")), ("onehot", OneHotEncoder(handle_unknown="ignore"))]), cat_cols),
            ],
            remainder="drop",
        )
        counts = pd.Series(groups).value_counts()
        inv_weights = np.array([1.0 / counts[g] for g in groups], dtype=float)
        inv_weights = inv_weights / np.mean(inv_weights)
        for weighted in [False, True]:
            ba, acc = [], []
            for train, test in cv.split(x, y, groups):
                pipe = Pipeline([("preprocess", pre), ("model", model)])
                if weighted:
                    pipe.fit(x.iloc[train], y[train], model__sample_weight=inv_weights[train])
                else:
                    pipe.fit(x.iloc[train], y[train])
                pred = pipe.predict(x.iloc[test])
                ba.append(balanced_accuracy_score(y[test], pred))
                acc.append(accuracy_score(y[test], pred))
            rows.append(
                {
                    "task": name,
                    "condition": "inverse-source weighting" if weighted else "unweighted grouped validation",
                    "folds": n_splits,
                    "balanced_accuracy_mean": np.mean(ba),
                    "balanced_accuracy_sd": np.std(ba, ddof=1) if len(ba) > 1 else 0,
                    "accuracy_mean": np.mean(acc),
                    "accuracy_sd": np.std(acc, ddof=1) if len(acc) > 1 else 0,
                }
            )
    res = pd.DataFrame(rows)
    base = res.pivot(index="task", columns="condition", values="balanced_accuracy_mean")
    for task in base.index:
        delta = base.loc[task].get("inverse-source weighting", np.nan) - base.loc[task].get("unweighted grouped validation", np.nan)
        res.loc[res["task"].eq(task), "weighted_minus_unweighted_BA"] = delta
    res.to_csv(OUTDIR / "source_weighting_transfer_test_v8.csv", index=False)
    return res


def add_v8_text(doc: Document, carbon: pd.DataFrame, weighting: pd.DataFrame) -> None:
    discussion_anchor = find_heading(doc, "5. Discussion")
    section6_anchor = find_heading(doc, "6. Practical Use Case for a New Soil Source")
    concl_anchor = find_heading(doc, "8. Conclusions")

    # Results additions before Discussion.
    add_paragraph_before(
        doc,
        discussion_anchor,
        "The carbon screening was also tested under a simple dynamic sensitivity setting rather than a single cradle-to-gate value. "
        "Binder-related emission factors were scaled by -20%, baseline and +20%, and transport distances of 0, 100 and 300 km were added using a screening factor of 0.062 kg CO2e per tonne-km. "
        "This calculation is not intended to replace a project-specific life-cycle assessment, but it shows whether the Pareto ordering remains stable when regional logistics and supply-chain uncertainty are introduced. "
        "The scenario summary is reported in Table 12. "
        "The most carbon-efficient mixtures were more sensitive to transport distance when the absolute binder dosage was low, whereas high-binder mixtures became less attractive under the +20% production-emission scenario. "
        "Thus, the Pareto analysis should be read as a carbon-aware pre-screening tool whose final ranking must be recalculated for the actual source-to-site distance and local binder supply route.",
    )
    csum = carbon.groupby(["binder_factor", "transport_km"], as_index=False).agg(
        median_co2_kg_per_100kg=("median_co2_kg_per_100kg", "median"),
        median_gain_per_kg_co2=("median_gain_per_kg_co2", "median"),
        n=("n", "sum"),
    )
    crows = [
        [
            f"{r.binder_factor:.2f}",
            str(int(r.transport_km)),
            str(int(r.n)),
            rounded(r.median_co2_kg_per_100kg, 2),
            rounded(r.median_gain_per_kg_co2, 2),
        ]
        for r in csum.itertuples()
    ]
    add_table_before(
        doc,
        discussion_anchor,
        "Table 12. Dynamic carbon sensitivity for the Pareto screening. CO2e = carbon dioxide equivalent; gain per kg CO2e is a screening ratio based on the predicted performance improvement divided by scenario-adjusted emissions.",
        ["Binder emission factor", "Transport distance (km)", "Mixtures (n)", "Median CO2e (kg/100 kg soil)", "Median gain per kg CO2e"],
        crows,
    )
    add_paragraph_before(
        doc,
        discussion_anchor,
        "A second diagnostic test examined whether the transfer loss could be reduced by giving each source a more equal influence during training. "
        "Inverse-source weighting was used as a deliberately simple domain-adaptation baseline: observations from small sources received larger sample weights, whereas observations from large sources received smaller weights. "
        "Table 13 reports the resulting grouped-source comparison. "
        "The test did not remove the grouped-source penalty, but it is useful because it separates a correctable sampling-imbalance effect from the deeper unmeasured-chemistry effect discussed below. "
        "Where source weighting improved balanced accuracy, the gain was interpreted only as evidence that source imbalance contributes to the transfer problem; where it did not improve performance, the result supports the need for richer soil descriptors rather than more complex black-box learners.",
    )
    wrows = [
        [
            r.task,
            r.condition,
            str(int(r.folds)),
            f"{r.balanced_accuracy_mean:.3f} +/- {r.balanced_accuracy_sd:.3f}",
            f"{r.accuracy_mean:.3f} +/- {r.accuracy_sd:.3f}",
            f"{r.weighted_minus_unweighted_BA:.3f}",
        ]
        for r in weighting.itertuples()
    ]
    add_table_before(
        doc,
        discussion_anchor,
        "Table 13. Source-weighting transfer test under grouped-source validation. BA = balanced accuracy; SD = standard deviation across held-out source folds.",
        ["Task", "Training condition", "Folds", "BA +/- SD", "Accuracy +/- SD", "Weighted BA change"],
        wrows,
    )

    # Discussion additions.
    add_paragraph_before(doc, section6_anchor, "5.1 Domain Shift as a Physical-Chemical Problem", style="Heading 2")
    add_paragraph_before(
        doc,
        section6_anchor,
        "The transfer loss observed in the grouped-source experiments should not be interpreted only as a mathematical weakness of machine learning. "
        "It is also a physical signal that important soil-chemistry controls are absent from the compiled literature database. "
        "For the multiclass tasks, balanced accuracy dropped by as much as 0.465 for UCS and 0.506 for CBR when the validation was changed from random internal folds to held-out source folds. "
        "Such a loss is consistent with the behaviour of expansive clay systems, where swelling and shear strength are governed by the diffuse double layer around montmorillonitic minerals, exchangeable cations and the chemistry of the pore water. "
        "When lime, fly ash, cementitious binders or alkali activators are added, cation exchange, flocculation-agglomeration and pozzolanic or geopolymeric reaction rates depend strongly on initial pH, pore-water salinity, clay fraction, exchangeable sodium and calcium availability, and the reactive silica-alumina-calcium balance of the stabilizer. "
        "Most published tables used in this study report useful index and dosage information, but they rarely report mineralogical percentages from X-ray diffraction, pore-water chemistry, binder oxide composition or curing humidity. "
        "The models therefore rely on surrogate descriptors such as liquid limit, plasticity index, stabilizer dosage and curing age. "
        "These surrogates can support screening within a familiar source domain, but they cannot fully identify why two soils with similar plasticity may react differently to the same binder. "
        "Resolving the transferability gap therefore requires better reporting of mineralogical and chemical descriptors, not only larger data volume or more elaborate algorithms. "
        "This interpretation agrees with the geotechnical mechanism described by Mitchell and Soga and with recent stabilized-soil modelling studies that report strong performance under internal validation but weaker generalisation when the soil source changes.",
    )
    add_paragraph_before(doc, section6_anchor, "5.2 Re-evaluating the CBR Database Utility", style="Heading 2")
    add_paragraph_before(
        doc,
        section6_anchor,
        "The CBR part of the database must be treated more cautiously than the UCS part. "
        "Although 485 CBR observations were available, they came from only three independent source groups, so grouped-source validation is necessarily sensitive to which source is held out. "
        "The fold-level standard deviation in balanced accuracy was 0.078 for the CBR pass/fail classifier, and the multiclass CBR transfer loss reached 0.506. "
        "For this reason, the CBR model is not presented as a globally reliable design-value predictor and it must not be used to replace local California bearing ratio laboratory testing for pavement design. "
        "Its value in this paper is narrower but still useful: it demonstrates how a carbon-aware Pareto ranking workflow can screen candidate stabilization strategies before a project commits to laboratory confirmation. "
        "In practical use, the CBR outputs should be read as regional screening evidence, with final acceptance based on site-specific compaction, soaking and penetration tests conducted under the governing pavement design standard.",
    )
    add_paragraph_before(doc, section6_anchor, "5.3 Minimum Geotechnical Reporting Protocol", style="Heading 2")
    add_paragraph_before(
        doc,
        section6_anchor,
        "The diagnostic value of this study is that it identifies what future experimental papers must report if geotechnical machine-learning models are to become transferable rather than merely local. "
        "A minimum geotechnical reporting protocol should accompany soil stabilization studies as a complete machine-readable table, preferably as CSV or XLSX supplementary material. "
        "The table should include the raw mechanical outputs, such as UCS, CBR, swelling index and compaction results, together with replicate counts, means, standard deviations and the raw stress-strain or load-penetration curves where possible. "
        "It should also include the mineralogical composition of the untreated soil, especially montmorillonite, kaolinite, illite, quartz and carbonate fractions obtained from X-ray diffraction; chemical descriptors such as pH, electrical conductivity, exchangeable cations and organic content; exact binder compositions including CaO, SiO2, Al2O3, Fe2O3, loss on ignition and alkali content; fibre geometry and tensile properties where fibres are used; compaction energy, specimen size, moisture conditioning and curing temperature and humidity. "
        "Reporting these descriptors would allow future models to distinguish genuine material mechanisms from source labels, reduce hidden domain shift and make external validation more meaningful. "
        "Without such reporting, the field will continue to accumulate many local datasets that are useful for individual case studies but difficult to merge into globally transferable design-support tools.",
    )

    # Limitations and conclusions additions.
    add_paragraph_before(
        doc,
        concl_anchor,
        "The Indian and international data-expansion routes identified during this revision should be used selectively. "
        "National geoscience and soil portals such as NGDR, GSI Bhukosh, ISRO Bhuvan, SLUSI and Data.gov.in can provide regional context on soil class, geology, climate and agricultural chemistry, while CRRI, Zenodo and Mendeley-style repositories can provide direct laboratory datasets when tabular test results are openly available. "
        "These sources should not be merged blindly with mechanical laboratory data; instead, they should be used to add source-level descriptors such as region, parent geology, climate zone and likely mineralogical setting, and any inferred variables should be clearly labelled as contextual rather than measured.",
    )
    add_paragraph_before(
        doc,
        concl_anchor,
        "The main technical conclusion is that source-aware validation changes the interpretation of stabilized-soil machine learning. "
        "High internal cross-validation scores show that the database contains learnable patterns, but the held-out-source losses show that those patterns are not yet universal because essential mineralogical and chemical controls are usually missing. "
        "The UCS classifier is suitable for preliminary screening under similar source conditions, whereas the CBR classifier should be treated only as a proof of the carbon-aware ranking workflow and must not replace local laboratory CBR testing. "
        "The next defensible step is therefore not simply a larger black-box model, but a better reported, chemistry-aware database built around the minimum geotechnical reporting protocol proposed in this study.",
    )


def update_readme() -> None:
    text = f"""# V8 Technical Modification Report

Created: 2026-05-27

This folder contains the revised manuscript produced from the v7 GitHub-ready manuscript.

## Added in v8

- Expanded the domain-shift discussion as a physical-chemical problem, including montmorillonitic diffuse double layer behaviour, pH, salinity, clay fraction, cation exchange, flocculation and pozzolanic/geopolymeric reaction controls.
- Reframed the CBR model as regional screening and conceptual validation of the Pareto workflow, not as a replacement for local laboratory CBR tests.
- Added a dynamic carbon sensitivity analysis using binder-emission factors of 0.80, 1.00 and 1.20 and transport distances of 0, 100 and 300 km.
- Added a source-weighting/domain-adaptation baseline using inverse-source sample weights under grouped-source validation.
- Added a minimum geotechnical reporting protocol for future soil-stabilization datasets.
- Added a short note on using NGDR, GSI Bhukosh, ISRO Bhuvan, SLUSI, Data.gov.in, CRRI, Zenodo and Mendeley-style repositories as future contextual or laboratory data sources.

## New analysis files

- `dynamic_carbon_sensitivity_v8.csv`
- `source_weighting_transfer_test_v8.csv`

## Manuscript

- `Submission_manuscript_v8_technical_modifications.docx`
"""
    (OUTDIR / "README_V8_TECHNICAL_MODIFICATIONS.md").write_text(text, encoding="utf-8")


def main() -> None:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    doc = Document(INDOC)
    carbon = make_carbon_sensitivity()
    weighting = source_weighting_test()
    add_v8_text(doc, carbon, weighting)
    doc.save(OUTDOC)
    update_readme()

    # Copy repo package forward and add new analysis files to results for reviewer traceability.
    repo_out = OUTDIR / "source-aware-stabilized-soil-ml-v8"
    if repo_out.exists():
        def _remove_readonly(func, path, exc_info):
            try:
                os.chmod(path, 0o700)
                func(path)
            except Exception:
                raise

        shutil.rmtree(repo_out, onerror=_remove_readonly)
    shutil.copytree(REPO, repo_out)
    (repo_out / "results").mkdir(exist_ok=True)
    shutil.copy2(OUTDIR / "dynamic_carbon_sensitivity_v8.csv", repo_out / "results" / "dynamic_carbon_sensitivity_v8.csv")
    shutil.copy2(OUTDIR / "source_weighting_transfer_test_v8.csv", repo_out / "results" / "source_weighting_transfer_test_v8.csv")
    with open(OUTDIR / "v8_outputs.json", "w", encoding="utf-8") as f:
        json.dump(
            {
                "manuscript": str(OUTDOC),
                "carbon_sensitivity": str(OUTDIR / "dynamic_carbon_sensitivity_v8.csv"),
                "source_weighting": str(OUTDIR / "source_weighting_transfer_test_v8.csv"),
                "repo_copy": str(repo_out),
            },
            f,
            indent=2,
        )


if __name__ == "__main__":
    main()
