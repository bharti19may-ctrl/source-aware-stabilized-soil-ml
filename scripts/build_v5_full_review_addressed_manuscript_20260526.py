from pathlib import Path
import re
import shutil
import textwrap
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT = Path(r"D:\kec folder\my files\student project\Elsevier_Q2_ExternalOnly_StabilizedSoil_ML")
DATA = PROJECT / "Final_Reliable_Dataset_For_Elsevier_v1_2026-05-25"
DEC = PROJECT / "Decision_Support_Framework_v1_2026-05-25"
RG = PROJECT / "Reviewer_Grade_Figures_Tables_v1_2026-05-26"
REQ = PROJECT / "Review_Report_Required_Analyses_v1_2026-05-26"
SYN = PROJECT / "Synthetic_Augmentation_PrePaper_Audit_v1_2026-05-26"
OUT = PROJECT / "Submission_Manuscript_v5_FullReviewAddressed_2026-05-26"
FIG = OUT / "figures"
TAB = OUT / "tables"
for p in [OUT, FIG, TAB]:
    p.mkdir(parents=True, exist_ok=True)


def load():
    d = {
        "ucs": pd.read_csv(DATA / "final_reliable_UCS_dataset_v1.csv"),
        "cbr": pd.read_csv(DATA / "final_reliable_CBR_dataset_v1.csv"),
        "source": pd.read_csv(RG / "tables" / "table_data_source_inventory.csv"),
        "class_best": pd.read_csv(DEC / "best_decision_support_classification_v1.csv"),
        "ranking": pd.read_csv(DEC / "decision_support_ranking_metrics_v1.csv"),
        "synth": pd.read_csv(SYN / "synthetic_classification_vs_real_only_v1.csv"),
        "reg": pd.read_csv(REQ / "tables" / "table_regression_baseline_and_augmented_comparison.csv"),
        "uncertainty": pd.read_csv(REQ / "tables" / "table_fold_uncertainty_classification.csv"),
        "threshold": pd.read_csv(REQ / "tables" / "table_threshold_sensitivity_class_balance.csv"),
        "hyper": pd.read_csv(REQ / "tables" / "table_model_hyperparameters_reproducibility.csv"),
        "refs": pd.read_csv(RG / "tables" / "expanded_reference_backbone_50_items.csv"),
        "pareto": pd.read_csv(DEC / "pareto_front_real_mixtures_v1.csv"),
    }
    d["carbon"] = make_carbon_sensitivity(d["pareto"])
    return d


def make_carbon_sensitivity(pareto):
    df = pareto.copy()
    if "kgco2e_per_100kg_soil" not in df.columns or "gain" not in df.columns:
        return pd.DataFrame(columns=["response", "carbon_factor_multiplier", "candidate_count", "median_gain", "median_co2", "median_gain_per_co2", "top10_mean_gain_per_co2"])
    df["kgco2e_per_100kg_soil"] = pd.to_numeric(df["kgco2e_per_100kg_soil"], errors="coerce")
    df["gain"] = pd.to_numeric(df["gain"], errors="coerce")
    df = df[df["kgco2e_per_100kg_soil"].notna() & df["gain"].notna() & (df["kgco2e_per_100kg_soil"] > 0)].copy()
    rows = []
    for mult in [0.8, 1.0, 1.2]:
        temp = df.copy()
        temp["scenario_co2"] = temp["kgco2e_per_100kg_soil"] * mult
        temp["scenario_gain_per_co2"] = temp["gain"] / temp["scenario_co2"]
        for response, g in temp.groupby("pareto_response"):
            top = g.sort_values("scenario_gain_per_co2", ascending=False).head(10)
            rows.append({
                "response": response,
                "carbon_factor_multiplier": mult,
                "candidate_count": len(g),
                "median_gain": g["gain"].median(),
                "median_co2": g["scenario_co2"].median(),
                "median_gain_per_co2": g["scenario_gain_per_co2"].median(),
                "top10_mean_gain_per_co2": top["scenario_gain_per_co2"].mean(),
            })
    return pd.DataFrame(rows)


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[name]
        st.font.name = "Arial"
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.size = Pt(11 if name == "Normal" else 18 if name == "Title" else 15 if name == "Heading 1" else 12)


def para(doc, text):
    for block in textwrap.dedent(text).strip().split("\n\n"):
        doc.add_paragraph(block.strip())


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def number(doc, text):
    doc.add_paragraph(text, style="List Number")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, val in [("top", 80), ("start", 120), ("bottom", 80), ("end", 120)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def add_table(doc, df, caption, font_size=7, max_rows=None):
    p = doc.add_paragraph()
    p.add_run(caption).bold = True
    if max_rows is not None:
        df = df.head(max_rows)
    df = df.copy()
    for c in df.columns:
        if pd.api.types.is_float_dtype(df[c]):
            df[c] = df[c].map(lambda x: "" if pd.isna(x) else f"{x:.3g}")
    df = df.fillna("").astype(str)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = c
        shade(cell, "E8EEF5")
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(font_size)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(df.columns):
            cells[j].text = str(row[c])
            margins(cells[j])
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for par in cells[j].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph()


def add_fig(doc, path, caption, width=6.3):
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    for r in p.runs:
        r.font.size = Pt(9)


def equation(doc, text, num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{text}     ({num})")
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)


def source_codes(source):
    s = source.copy().sort_values("total_rows", ascending=False).reset_index(drop=True)
    s["Source code"] = [f"S{i+1}" for i in range(len(s))]
    return s


def make_figures(d):
    plt.rcParams.update({"font.size": 10, "axes.spines.top": False, "axes.spines.right": False})
    paths = {}
    src = source_codes(d["source"])
    src.to_csv(TAB / "table_source_code_key.csv", index=False)

    # Figure 1: clean workflow, no tiny text.
    fig, ax = plt.subplots(figsize=(9.4, 5.6))
    ax.axis("off")
    steps = [
        ("1", "Traceable literature records", "UCS and CBR records with source identity"),
        ("2", "Data curation", "Unit check, plausibility screening, harmonisation"),
        ("3", "Real-only modelling set", "Synthetic records excluded from validation data"),
        ("4", "Validation design", "Internal k-fold and grouped-source testing"),
        ("5", "Decision modelling", "Classification, ranking, feature importance"),
        ("6", "Engineering screening", "Augmentation sensitivity and carbon Pareto front"),
    ]
    ys = np.linspace(0.88, 0.12, len(steps))
    for (n, title, subtitle), y in zip(steps, ys):
        ax.text(0.08, y, n, ha="center", va="center", fontsize=13, fontweight="bold",
                bbox=dict(boxstyle="circle,pad=0.35", fc="#315C7A", ec="#315C7A"), color="white")
        ax.text(0.16, y + 0.025, title, ha="left", va="center", fontsize=11, fontweight="bold")
        ax.text(0.16, y - 0.035, subtitle, ha="left", va="center", fontsize=9.5, color="#444444")
        if y != ys[-1]:
            ax.annotate("", xy=(0.08, y - 0.10), xytext=(0.08, y - 0.045),
                        arrowprops=dict(arrowstyle="->", lw=1.2, color="#555555"))
    ax.text(0.62, 0.12, "Validation rule: generated samples never enter test folds.",
            ha="center", va="center", fontsize=10, color="#7A1F3D",
            bbox=dict(boxstyle="round,pad=0.4", fc="#F7E8ED", ec="#7A1F3D"))
    fig.tight_layout()
    p = FIG / "fig1_clear_methodological_workflow.png"
    fig.savefig(p, dpi=300)
    plt.close(fig)
    paths["workflow"] = p

    # Figure 2: source distribution with source codes.
    plot = src.sort_values("total_rows", ascending=True)
    fig, ax = plt.subplots(figsize=(7.4, 4.8))
    ax.barh(plot["Source code"], plot["ucs_rows"], label="UCS", color="#315C7A")
    ax.barh(plot["Source code"], plot["cbr_rows"], left=plot["ucs_rows"], label="CBR", color="#8B1E3F")
    ax.set_xlabel("Number of real observations")
    ax.set_ylabel("Literature source code")
    ax.set_title("Dataset Contribution by Source")
    ax.legend(frameon=False)
    fig.tight_layout()
    p = FIG / "fig2_clear_source_distribution.png"
    fig.savefig(p, dpi=300)
    plt.close(fig)
    paths["source_distribution"] = p

    # Figure 3: availability by source code.
    avail = pd.read_csv(RG / "tables" / "table_feature_group_availability_by_source.csv")
    avail = avail.merge(src[["source_dataset", "Source code"]], on="source_dataset", how="left")
    heat = avail.set_index("Source code").drop(columns=["source_dataset"])
    fig, ax = plt.subplots(figsize=(7.8, 4.8))
    im = ax.imshow(heat.values, aspect="auto", cmap="YlGnBu", vmin=0, vmax=1)
    ax.set_xticks(np.arange(len(heat.columns)))
    ax.set_xticklabels(heat.columns, rotation=30, ha="right")
    ax.set_yticks(np.arange(len(heat.index)))
    ax.set_yticklabels(heat.index)
    ax.set_title("Feature Availability by Source")
    cbar = fig.colorbar(im, ax=ax)
    cbar.set_label("Non-missing fraction")
    fig.tight_layout()
    p = FIG / "fig3_clear_feature_availability.png"
    fig.savefig(p, dpi=300)
    plt.close(fig)
    paths["feature_availability"] = p

    # Figure 4 validation comparison.
    class_best = d["class_best"].copy()
    class_best["Task"] = class_best["dataset"].str.replace("_kfold", "", regex=False).str.replace("_", " ")
    fig, ax = plt.subplots(figsize=(8.0, 4.6))
    for validation, g in class_best.groupby("validation"):
        g = g.sort_values("Task")
        ax.plot(g["Task"], g["balanced_accuracy"], marker="o", label=validation)
    ax.set_ylabel("Balanced accuracy")
    ax.set_ylim(0, 1)
    ax.set_title("Internal Validation Compared with Source-Grouped Transfer")
    ax.tick_params(axis="x", rotation=25)
    ax.legend(fontsize=8, frameon=False)
    fig.tight_layout()
    p = FIG / "fig4_validation_transfer_gap.png"
    fig.savefig(p, dpi=300)
    plt.close(fig)
    paths["validation_gap"] = p

    # Copy established figures.
    copy_map = {
        "ucs_confusion": "fig_confusion_UCS_high_binary_group_by_source_5fold_ExtraTrees.png",
        "cbr_confusion": "fig_confusion_CBR_pass_8_binary_group_by_source_3fold_RF.png",
        "ucs_importance": "fig_feature_importance_ucs_high_binary.png",
        "cbr_importance": "fig_feature_importance_cbr_pass_8.png",
        "synthetic": "fig_synthetic_ratio_sensitivity.png",
        "pareto": "fig_real_pareto_carbon_strength.png",
    }
    for key, name in copy_map.items():
        srcp = RG / "figures" / name
        dst = FIG / name
        shutil.copyfile(srcp, dst)
        paths[key] = dst
    for name in ["fig_threshold_sensitivity_class_balance.png", "fig_carbon_factor_sensitivity.png"]:
        srcp = REQ / "figures" / name
        if name == "fig_carbon_factor_sensitivity.png" and not srcp.exists():
            # Rebuild the carbon sensitivity figure locally when the earlier package did not export it.
            carbon = make_carbon_sensitivity(pd.read_csv(DEC / "pareto_front_real_mixtures_v1.csv"))
            fig, ax = plt.subplots(figsize=(7.2, 4.5))
            for response, g in carbon.groupby("response"):
                ax.plot(g["carbon_factor_multiplier"], g["top10_mean_gain_per_co2"], marker="o", label=response)
            ax.set_xlabel("Emission-factor multiplier")
            ax.set_ylabel("Top-10 mean gain per kg CO2e")
            ax.set_title("Carbon-Factor Sensitivity of Pareto-Efficient Mixtures")
            ax.legend(frameon=False)
            fig.tight_layout()
            dst = FIG / name
            fig.savefig(dst, dpi=300)
            plt.close(fig)
            paths[name] = dst
        elif srcp.exists():
            dst = FIG / name
            shutil.copyfile(srcp, dst)
            paths[name] = dst
    return paths, src


def format_tables(d, src):
    lit = pd.DataFrame([
        ["Stabilization mechanisms", "Fly ash, lime, cementitious binders, fibres, alkali-activated binders", "Explains expected variables and physical controls"],
        ["ML prediction studies", "UCS, CBR, swelling, compaction, XAI and physics-informed models", "Shows why high random-split accuracy alone is not sufficient"],
        ["Data augmentation", "GAN, ACGAN, physics/domain constrained synthetic samples", "Supports training-only sensitivity analysis"],
        ["Sustainability screening", "Carbon cost, binder intensity, Pareto ranking", "Supports performance-carbon mixture screening"],
    ], columns=["Theme", "Representative literature", "Use in this study"])
    lit.to_csv(TAB / "table1_literature_positioning.csv", index=False)

    source = src[["Source code", "source_dataset", "domain_family", "total_rows", "ucs_rows", "cbr_rows", "variables_available"]].copy()
    source.columns = ["Code", "Source", "Domain", "Rows", "UCS", "CBR", "Variables"]
    source.to_csv(TAB / "table2_source_inventory.csv", index=False)

    hyper = d["hyper"].copy()
    hyper.to_csv(TAB / "table3_model_configuration.csv", index=False)
    threshold = d["threshold"].copy()
    threshold.to_csv(TAB / "table4_threshold_sensitivity.csv", index=False)

    class_tbl = d["class_best"][["dataset", "validation", "model", "balanced_accuracy", "accuracy", "weighted_f1", "rows", "sources"]].copy()
    class_tbl.columns = ["Task", "Validation", "Model", "BA", "Accuracy", "Weighted F1", "Rows", "Sources"]
    class_tbl.to_csv(TAB / "table5_classification_results.csv", index=False)

    unc = d["uncertainty"].copy()
    keep = (
        (unc["dataset"].isin(["UCS_high_binary", "CBR_pass_8_binary", "UCS_multiclass"])) &
        (unc["validation"].str.contains("group|kfold", regex=True))
    )
    unc = unc[keep].sort_values(["dataset", "validation", "balanced_accuracy_mean"], ascending=[True, True, False]).groupby(["dataset", "validation"]).head(1)
    unc.to_csv(TAB / "table6_uncertainty.csv", index=False)

    reg = d["reg"][["dataset", "validation", "case", "model", "augmentation_multiplier", "r2", "rmse", "mae", "delta_r2_vs_same_model_baseline"]].copy()
    reg.to_csv(TAB / "table7_regression_benchmark.csv", index=False)

    synth = d["synth"]
    grouped = synth[(synth["validation"].str.contains("group")) & (synth["augmentation_multiplier"] > 0)].copy()
    synth_best = grouped.sort_values(["dataset", "label_col", "delta_balanced_accuracy"], ascending=[True, True, False]).groupby(["dataset", "label_col"]).head(1)
    synth_tbl = synth_best[["dataset", "validation", "model", "augmentation_multiplier", "balanced_accuracy_baseline", "balanced_accuracy", "delta_balanced_accuracy", "accuracy", "weighted_f1"]].copy()
    synth_tbl.columns = ["Task", "Validation", "Model", "Ratio", "Real-only BA", "Augmented BA", "Delta BA", "Accuracy", "Weighted F1"]
    synth_tbl.to_csv(TAB / "table8_augmentation_sensitivity.csv", index=False)

    rank = d["ranking"].copy()
    rank.columns = ["Dataset", "Target", "Objective", "Rows", "Sources", "Mean Spearman", "Median Spearman", "Mean Kendall", "Mean top-20 hit", "Median top-20 hit"]
    rank.to_csv(TAB / "table9_ranking.csv", index=False)

    carbon = d["carbon"].copy()
    carbon.to_csv(TAB / "table10_carbon_sensitivity.csv", index=False)
    stat = pd.DataFrame([
        ["Internal vs grouped-source validation", "Effect-size comparison using balanced accuracy", "Large drop for all main tasks", "Supports the transferability gap claim"],
        ["Fold variability", "Mean and standard deviation across folds", "Grouped-source folds show larger variability", "Avoids overclaiming single metric values"],
        ["Synthetic augmentation", "Real-only vs augmented-training effect size", "CBR improves modestly; UCS changes little", "Augmentation retained only as sensitivity analysis"],
        ["Formal paired significance test", "Not used as a primary decision criterion", "Grouped-source folds are only 3-5 independent sources", "Effect size and uncertainty are more appropriate than p-values for this dataset"],
    ], columns=["Comparison", "Statistical treatment", "Observed result", "Interpretation"])
    stat.to_csv(TAB / "table11_statistical_comparison_strategy.csv", index=False)
    return lit, source, hyper, threshold, class_tbl, unc, reg, synth_tbl, rank, carbon, stat


def cite_audit(doc_text, refs):
    rows = []
    for ref in refs["reference"].tolist():
        m = re.match(r"([^,]+)", ref)
        first = m.group(1) if m else ref.split()[0]
        year = re.search(r"(19|20)\d{2}", ref)
        y = year.group(0) if year else ""
        cited = (first in doc_text and y in doc_text) or (first.split()[0] in doc_text and y in doc_text)
        rows.append({"reference": ref, "first_author_or_standard": first, "year": y, "likely_cited": cited})
    out = pd.DataFrame(rows)
    out.to_csv(TAB / "reference_citation_audit.csv", index=False)
    return out


def build_doc():
    d = load()
    paths, src = make_figures(d)
    tables = format_tables(d, src)
    lit, source, hyper, threshold, class_tbl, unc, reg, synth_tbl, rank, carbon, stat = tables
    doc = Document()
    style_doc(doc)

    t = doc.add_paragraph(style="Title")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run("Source-aware machine learning for stabilized expansive soils: transfer validation, decision screening and carbon-efficient mixture ranking").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Author details to be inserted").italic = True

    doc.add_heading("Highlights", level=1)
    for h in [
        "A real-only literature dataset was curated for stabilized-soil UCS and CBR screening.",
        "Grouped-source validation quantified transfer loss across independent experimental sources.",
        "Model configuration, thresholds and uncertainty were reported for reproducibility.",
        "Synthetic augmentation gave limited CBR gains and was restricted to training folds.",
        "Carbon-aware Pareto screening produced practical shortlists for laboratory verification.",
    ]:
        bullet(doc, h)

    doc.add_heading("Abstract", level=1)
    para(doc, f"""
    Machine-learning models can assist stabilized-soil mixture selection, but their usefulness depends on transferability across independent soil sources and laboratory programmes. This study develops a source-aware framework using a curated real experimental dataset containing {len(d['ucs']):,} unconfined compressive strength (UCS) records and {len(d['cbr']):,} California bearing ratio (CBR) records. The workflow includes traceability screening, feature-availability assessment, internal and grouped-source validation, threshold sensitivity, model-configuration reporting, fold-level uncertainty, regression benchmarking, feature-importance analysis, training-only synthetic augmentation, source-wise ranking and carbon-aware Pareto screening. Internal validation produced high binary classification performance, but grouped-source validation was substantially weaker, showing that random splitting overestimates transferability for heterogeneous literature data. Continuous regression remained weak under grouped-source testing, particularly for UCS, supporting the shift toward classification and ranking. Synthetic augmentation modestly improved selected CBR screening metrics but did not create a universal prediction model. The most defensible use of the dataset is therefore preliminary mixture classification, prioritisation and carbon-efficient screening before local laboratory confirmation.
    """)
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("expansive soil; stabilized soil; machine learning; source-grouped validation; CBR; UCS; carbon efficiency; synthetic augmentation")

    doc.add_heading("1. Introduction", level=1)
    para(doc, """
    Expansive fine-grained soils are problematic in transportation infrastructure because moisture fluctuations can cause swelling, shrinkage, cracking and loss of support. Stabilization with lime, cement, fly ash, slag, rice husk ash, pond ash, alkali-activated binders and fibres has been widely investigated for improving subgrade and subbase performance (Katti 1978; Phani Kumar and Sharma 2004; Chauhan and Rajesh 2015; Dixit et al. 2016; Tiwari et al. 2016; Deshpande and Puranik 2017; Dixit 2017; Patil et al. 2018; Indiramma et al. 2020; Sharma and Kumar 2020; Ahmad et al. 2024; Parihar and Gupta 2024).

    The improvement mechanisms are well established in principle but variable in magnitude. Calcium-rich binders can induce cation exchange, flocculation, pozzolanic reaction and cementitious bonding. Fly ash and other aluminosilicate materials can contribute to secondary reaction products when calcium or alkaline activators are available (Luhar and Luhar 2022; Wang et al. 2023; Nanda and Priya 2024; Yang et al. 2025). Fibres can improve ductility and residual strength by bridging cracks and restraining tensile deformation (Rao et al. 2023; Navagire et al. 2025). However, response depends on soil mineralogy, plasticity, gradation, compaction condition, additive chemistry, fibre geometry, curing age and testing method.

    Machine learning has increasingly been used to model UCS, CBR, compaction and swelling behaviour in improved soils (Eyo and Onyekpe 2021; Eyo et al. 2022; Saad et al. 2023; Almuaythir et al. 2025; Huang et al. 2025; Thapa and Ghani 2024; Zeini et al. 2023). Many studies report high accuracy, but most use internal splits. When records from the same source appear in both training and testing subsets, the resulting accuracy can partly reflect interpolation within one experimental programme rather than transfer to a new source. Recent data-driven geotechnics discussions have therefore emphasised the need for better databases, transparent validation and open workflows (Ge et al. 2024; Li et al. 2025; Zong et al. 2025).
    """)

    doc.add_heading("2. Research Positioning and Gap", level=1)
    para(doc, """
    This study was designed around the limitation above. It does not attempt to present another high-accuracy black-box regression model. Instead, it asks which engineering decisions can be supported by heterogeneous literature data when independent-source transfer is evaluated explicitly. Table 1 summarises how the present work is positioned relative to the literature.

    The specific gap is threefold. First, literature-derived stabilized-soil datasets need source-aware validation because independent studies differ in materials and methods. Second, threshold-based classification requires justification because strength thresholds can appear arbitrary if class balance is not examined. Third, carbon-aware mixture screening needs to be connected to model outputs rather than added as a disconnected sustainability statement. The present framework addresses these points through grouped-source validation, threshold sensitivity, regression benchmarking, ranking and Pareto screening.
    """)
    add_table(doc, lit, "Table 1. Literature positioning and how each theme is used in the present study.", font_size=8)

    doc.add_heading("3. Data and Methods", level=1)
    doc.add_heading("3.1 Dataset Curation and Source Structure", level=2)
    para(doc, """
    The dataset was assembled from traceable published and open data sources. Records were retained only when the target value, source identity and essential mixture variables could be interpreted. Records with unresolved units, physically implausible values or uncertain target definitions were excluded. Standard geotechnical terminology and testing concepts follow relevant soil and pavement testing references where applicable (Bureau of Indian Standards 1970, 2016; ASTM D638 2014; Indian Roads Congress 2012).

    Table 2 gives the source inventory and code key used in the figures. Figure 1 shows the complete workflow. Figure 2 shows that data are unevenly distributed by source, and Figure 3 shows that feature reporting is incomplete and source dependent. These two figures were redesigned with source codes to avoid hidden text and to make the dataset structure legible.
    """)
    add_table(doc, source, "Table 2. Source inventory and source-code key for the real experimental dataset.", font_size=7)
    add_fig(doc, paths["workflow"], "Figure 1. Methodological workflow for source-aware stabilized-soil decision modelling.", width=6.1)
    add_fig(doc, paths["source_distribution"], "Figure 2. Real UCS and CBR observations contributed by each source code in Table 2.", width=6.0)
    add_fig(doc, paths["feature_availability"], "Figure 3. Feature-group availability by source code; values indicate the non-missing fraction.", width=6.1)

    doc.add_heading("3.2 Model Inputs, Preprocessing and Algorithms", level=2)
    para(doc, """
    The input variables included source/domain descriptors, soil class, gradation, Atterberg limits, specific gravity, optimum moisture content, maximum dry density, water content, dry density, binder dosages, fibre dosages, activator dosages, curing days and CBR condition where available. Numerical variables were median-imputed and robust-scaled. Categorical variables were most-frequent-imputed and one-hot encoded. This preprocessing was applied inside each training fold to avoid information leakage.

    Three tree-based model families were used because stabilized-soil response is nonlinear and contains interactions among binder content, curing age, compaction state and soil index properties. Random forest models provide robust bagged-tree baselines. Extremely randomized trees add stronger randomisation and were useful for feature-importance analysis. Histogram-gradient boosting was used for regression benchmarks and for equation-guided synthetic target assignment because it can model nonlinear response with controlled regularisation. Table 3 gives the implementation and hyperparameters used for reproducibility.
    """)
    add_table(doc, hyper, "Table 3. Model families, implementation and main hyperparameters.", font_size=7)

    doc.add_heading("3.3 Targets, Metrics and Threshold Sensitivity", level=2)
    para(doc, """
    UCS high-strength screening used 1500 kPa as the main threshold, and CBR pass/fail screening used 8%. These thresholds were selected as practical screening levels, not universal design specifications. To address threshold arbitrariness, additional UCS thresholds of 1000, 2000 and 2500 kPa and CBR thresholds of 5, 10 and 15% were examined. Table 4 and Figure 4 report the resulting class-balance sensitivity.
    """)
    equation(doc, "y_UCS = 1 if UCS >= 1500 kPa, otherwise 0", 1)
    equation(doc, "y_CBR = 1 if CBR >= 8%, otherwise 0", 2)
    equation(doc, "BA = (1/K) Σ[TP_k/(TP_k+FN_k)]", 3)
    equation(doc, "ρ_s = 1 - 6Σd_i²/[n(n²-1)]", 4)
    equation(doc, "CO₂e_mix = Σ(m_j EF_j)", 5)
    equation(doc, "Gain = R_mix - R_baseline", 6)
    equation(doc, "Carbon efficiency = Gain/CO₂e_mix", 7)
    add_table(doc, threshold, "Table 4. Threshold sensitivity and class balance for UCS and CBR screening.", font_size=7)
    add_fig(doc, REQ / "figures" / "fig_threshold_sensitivity_class_balance.png", "Figure 4. Class-balance sensitivity for alternative UCS and CBR thresholds.", width=5.8)

    doc.add_heading("3.4 Validation, Uncertainty and Synthetic Augmentation", level=2)
    para(doc, """
    Two validation schemes were used. Stratified k-fold validation measured internal predictive consistency in the pooled dataset. Grouped-source validation left complete sources out of training and tested on unseen sources. Fold-level mean and standard deviation were calculated for key classification metrics. Continuous regression was retained as a benchmark to show why exact prediction was not the primary claim.

    Synthetic samples were generated only within training folds using bounded perturbation and equation-guided target assignment inside the observed source/domain range. Synthetic rows were not allowed in validation folds. Synthetic-to-real training ratios of 0.5, 1.0 and 2.0 were compared with the real-only baseline.
    """)

    doc.add_heading("4. Results", level=1)
    doc.add_heading("4.1 Classification and Transfer Gap", level=2)
    para(doc, """
    Table 5 shows that internal validation substantially outperformed grouped-source validation. UCS high-strength and CBR pass/fail screening achieved strong internal balanced accuracy, but the grouped-source values were lower. Figure 5 visualises this transfer gap. Figures 6 and 7 show the grouped-source confusion matrices for the two binary decision tasks. These results confirm that the paper should not claim universal exact prediction.
    """)
    add_table(doc, class_tbl, "Table 5. Best classification performance under internal and grouped-source validation.", font_size=7)
    add_fig(doc, paths["validation_gap"], "Figure 5. Difference between internal validation and source-grouped transfer performance.", width=6.0)
    add_fig(doc, paths["ucs_confusion"], "Figure 6. Grouped-source confusion matrix for UCS high-strength screening.", width=4.8)
    add_fig(doc, paths["cbr_confusion"], "Figure 7. Grouped-source confusion matrix for CBR pass/fail screening.", width=4.8)

    doc.add_heading("4.2 Uncertainty and Regression Benchmark", level=2)
    para(doc, """
    Table 6 reports fold-level uncertainty. Variability was larger for grouped-source validation because each fold represents a different source, not a random subset. Table 7 retains the continuous regression benchmark. Grouped-source regression remained weak, especially for UCS, even when augmentation was used. This directly supports the shift from point prediction to classification, ranking and screening.
    """)
    add_table(doc, unc, "Table 6. Fold-level uncertainty for principal classification tasks.", font_size=7)
    add_table(doc, reg, "Table 7. Continuous regression benchmark and training-only augmentation comparison.", font_size=7)
    add_table(doc, stat, "Table 8. Statistical comparison strategy used to avoid overclaiming minor model differences.", font_size=7)

    doc.add_heading("4.3 Feature Importance and Mechanistic Consistency", level=2)
    para(doc, """
    Figures 8 and 9 show global feature importance for UCS and CBR screening. The influential variables are consistent with expected geotechnical mechanisms: curing period and binder variables affect cementitious bonding; compaction descriptors affect density and moisture state; plasticity and gradation influence soil reactivity and bearing behaviour; CBR condition affects measured bearing capacity. The analysis therefore gives a physically interpretable model, although it remains limited by incomplete reporting across sources.
    """)
    add_fig(doc, paths["ucs_importance"], "Figure 8. Global feature importance for UCS high-strength classification.", width=6.0)
    add_fig(doc, paths["cbr_importance"], "Figure 9. Global feature importance for CBR pass/fail classification.", width=6.0)

    doc.add_heading("4.4 Augmentation, Ranking and Carbon Screening", level=2)
    para(doc, """
    Table 8 and Figure 10 show that synthetic augmentation produced modest and task-dependent gains. CBR pass/fail grouped-source balanced accuracy improved at a 0.5 augmentation ratio, whereas UCS high-strength classification changed very little. Table 9 shows that source-wise ranking was more stable than exact regression, with CBR performance ranking giving the strongest mean Spearman coefficient. Figure 11 and Table 10 present carbon-aware screening and sensitivity. The carbon ranking changed in magnitude when emission factors were scaled, but the decision logic remained stable: mixtures should be shortlisted by both performance gain and estimated carbon burden.
    """)
    add_table(doc, synth_tbl, "Table 9. Best grouped-source effect of training-only synthetic augmentation.", font_size=7)
    add_fig(doc, paths["synthetic"], "Figure 10. Synthetic-to-real training ratio sensitivity evaluated only on real validation records.", width=6.0)
    add_table(doc, rank, "Table 10. Source-wise mixture ranking performance.", font_size=7)
    add_fig(doc, paths["pareto"], "Figure 11. Real Pareto-efficient mixtures in performance-gain and estimated-carbon space.", width=6.0)
    add_table(doc, carbon, "Table 11. Carbon-factor sensitivity for Pareto-efficient mixtures.", font_size=7)
    add_fig(doc, paths["fig_carbon_factor_sensitivity.png"], "Figure 12. Sensitivity of carbon-efficiency ranking to emission-factor uncertainty.", width=5.8)

    doc.add_heading("5. Discussion", level=1)
    para(doc, """
    The main result is that stabilized-soil machine learning is highly sensitive to validation design. Internal validation indicates that the curated variables contain useful information, but grouped-source validation shows that those relationships are not fully transferable. The difference is expected because source studies differ in soil mineralogy, compaction energy, curing conditions, binder chemistry, fibre geometry and test procedures. It is also consistent with broader concerns in geotechnical data science about small, heterogeneous and incompletely reported datasets.

    The regression benchmark is important. If regression had been omitted entirely, the classification focus could appear arbitrary. Table 7 shows that grouped-source regression was retained and evaluated, but the results were not strong enough to support a universal numerical prediction model. The classification and ranking tasks are therefore not a retreat from analysis; they are a better match to the reliability of the available data and to preliminary engineering screening.

    The threshold sensitivity analysis reduces the risk that the binary classes appear arbitrary. The adopted thresholds retain usable class sizes and engineering meaning. Nevertheless, project-specific thresholds should be selected according to local specifications, pavement layer function and reliability requirements. The model outputs should be treated as screening indicators rather than design values.

    Synthetic augmentation has a limited role. It improved selected CBR tasks because CBR data were sparse, but it did not solve source-transfer limitations. Generated rows can regularise training but cannot replace missing laboratory evidence. The strict training-only rule is therefore essential.

    The carbon analysis adds practical value because strength alone is not enough for sustainable stabilization. A mixture with the highest UCS may not be preferable if a lower-carbon mixture satisfies the required class. The Pareto and sensitivity analyses provide a transparent shortlist for confirmatory laboratory testing.
    """)

    doc.add_heading("6. Practical Use Case for a New Soil Source", level=1)
    para(doc, """
    For a new expansive soil source, the framework should be used as a screening workflow rather than a direct design tool. First, the user enters the available index, compaction and mixture variables and checks whether the new soil falls inside the feature ranges of the curated database. Second, the binary classifiers estimate whether candidate mixtures are likely to reach the UCS or CBR screening class. Third, candidate mixtures are ranked within the closest source/domain group. Fourth, carbon efficiency is calculated to remove mixtures that require high binder-related emissions for only small performance gains. Finally, the highest-ranked and Pareto-efficient mixtures are selected for local compaction, UCS, CBR, swelling and durability testing.

    This use case clarifies the intended engineering role of the model. The model narrows the laboratory search space; it does not replace site-specific geotechnical testing. This interpretation is consistent with the grouped-source results, which show that transfer is useful for screening but not strong enough for final design values.
    """)

    doc.add_heading("7. Limitations and Reproducibility", level=1)
    para(doc, """
    The study remains limited by the information available in source papers. Mineralogy, pore-fluid chemistry, curing humidity, compaction energy and microstructural observations were not consistently available. The CBR dataset is smaller than the UCS dataset, so CBR conclusions should be treated as screening-level. Carbon calculations are approximate and should not be interpreted as a full life-cycle assessment. The supplementary material should include the cleaned real dataset, feature dictionary, source inventory, fold assignments, code, prediction files, augmentation rules and carbon factors.
    """)

    doc.add_heading("8. Conclusions", level=1)
    for c in [
        "The curated dataset is adequate for preliminary decision support but not for universal exact prediction across all stabilized soils.",
        "Internal validation substantially overestimated transfer performance compared with grouped-source validation.",
        "UCS high-strength classification, CBR pass/fail screening and source-wise ranking are more defensible than exact grouped-source regression for the current data.",
        "Synthetic augmentation should remain a training-only sensitivity tool because it produced only modest, task-dependent improvements.",
        "Carbon-aware Pareto screening is a useful practical output because it identifies mixtures that balance performance gain and estimated carbon burden.",
    ]:
        number(doc, c)

    doc.add_heading("CRediT Authorship Contribution Statement", level=1)
    para(doc, "Author contributions should be completed by the authors before submission.")
    doc.add_heading("Declaration of Competing Interest", level=1)
    para(doc, "The authors should declare whether they have any competing financial interests or personal relationships that could have influenced the work.")
    doc.add_heading("Ethics Statement", level=1)
    para(doc, "This study used previously published literature data and did not involve human participants, animals or field interventions requiring ethical approval.")
    doc.add_heading("Data Availability", level=1)
    para(doc, "The curated real dataset, source inventory, feature dictionary, scripts, fold assignments, prediction files, augmentation rules and carbon-factor table should be submitted as supplementary material and deposited in a permanent repository such as Mendeley Data, Zenodo or an institutional repository before submission. Synthetic rows are labelled separately and are not part of the real experimental dataset.")
    doc.add_heading("Appendix A. Revision and Reproducibility Schedule", level=1)
    para(doc, "The following schedule is included as a project-management appendix for completing the remaining reproducibility and submission tasks.")
    gantt = pd.DataFrame([
        ["Dataset and code repository", "Prepare cleaned data, scripts, fold IDs and readme", "3 days"],
        ["Reference verification", "Check every reference and DOI against source PDFs", "3 days"],
        ["Supplementary material", "Prepare data dictionary, augmentation rules and carbon factors", "2 days"],
        ["Final figures", "Inspect all figures at journal resolution and revise labels", "2 days"],
        ["Language and formatting", "Final grammar, table captions and journal checklist", "2 days"],
        ["Submission package", "Cover letter, highlights, declarations and final upload", "1 day"],
    ], columns=["Task", "Action", "Estimated duration"])
    add_table(doc, gantt, "Table A1. Gantt-style revision schedule for completing the submission package.", font_size=8)
    doc.add_heading("References", level=1)
    for ref in d["refs"]["reference"].tolist():
        doc.add_paragraph(ref)

    out = OUT / "Submission_manuscript_v5_full_review_addressed.docx"
    doc.save(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    audit = cite_audit(text, d["refs"])
    (OUT / "README_V5_FULL_REVIEW_ADDRESSED.md").write_text(
        "This v5 manuscript was rebuilt from scratch to address the review report more completely. It uses clearer Figures 1-3, ordered tables, expanded ML methodology, threshold sensitivity, uncertainty, regression benchmark, augmentation sensitivity, carbon sensitivity and citation audit.\n",
        encoding="utf-8",
    )
    print(out)
    print(OUT)
    print("uncited_likely", int((~audit["likely_cited"]).sum()))


if __name__ == "__main__":
    build_doc()
